import streamlit as st
import docx
import os
import re
import requests
from io import BytesIO
from langchain_mistralai import MistralAIEmbeddings, ChatMistralAI
from langchain_groq import ChatGroq
from langchain_xai.chat_models import ChatXAI
from langchain_core.globals import set_llm_cache
from langchain_community.cache import InMemoryCache
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS
import openai
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
import vertexai
from vertexai.preview.vision_models import ImageGenerationModel


# --- Configuración de la Aplicación ---
st.set_page_config(
    page_title="Visor de Documentos de Prompt Engineering",
    page_icon=None,
    layout="wide",
)

# --- Rutas a los archivos ---
# Para Streamlit Cloud usamos URLs de GitHub, para local usamos rutas locales
if os.path.exists("logo.png"):
    LOGO_PATH = "logo.png"
else:
    LOGO_PATH = "https://github.com/JulianTorrest/Agente-IA/raw/main/logo.png"

# URLs de los documentos en GitHub
GITHUB_PROMPTS_URL = "https://github.com/JulianTorrest/Agente-IA/raw/main/Ingenier%C3%ADa%20de%20Prompts.docx"
GITHUB_MANUAL_URL = "https://github.com/JulianTorrest/Agente-IA/raw/main/Manual%20del%20Creador%20de%20Agentes%20Copilot%20365.docx"

# Documentos locales (si existen)
BASE_PATH = os.path.dirname(__file__)
PROMPTS_DOC_PATH = os.path.join(BASE_PATH, "Ingeniería de Prompts.docx")
MANUAL_DOC_PATH = os.path.join(BASE_PATH, "Manual del Creador de Agentes Copilot 365.docx")

@st.cache_data
def descargar_documento_desde_github(url, filename):
    """Descarga un documento .docx desde GitHub."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        st.error(f"Error descargando {filename} desde GitHub: {e}")
        return None

@st.cache_data
def leer_documento_docx(path):
    """Lee un archivo .docx desde una ruta local o GitHub y devuelve el texto."""
    # Si no existe localmente, intentar descargar desde GitHub
    if not os.path.exists(path):
        if "Ingeniería de Prompts" in path:
            github_url = GITHUB_PROMPTS_URL
            filename = "Ingeniería de Prompts.docx"
        elif "Manual del Creador" in path:
            github_url = GITHUB_MANUAL_URL
            filename = "Manual del Creador de Agentes Copilot 365.docx"
        else:
            return f"**Error:** No se pudo encontrar el archivo en la ruta especificada: `{path}`. Por favor, verifica que el archivo existe."
        
        # Descargar desde GitHub
        doc_content = descargar_documento_desde_github(github_url, filename)
        if doc_content is None:
            return f"**Error:** No se pudo descargar {filename} desde GitHub."
        
        try:
            doc = docx.Document(doc_content)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"**Error:** Ocurrió un problema al leer el archivo `{filename}`. Detalles: {e}"
    
    # Si existe localmente, leerlo directamente
    try:
        doc = docx.Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"**Error:** Ocurrió un problema al leer el archivo `{path}`. Detalles: {e}"

@st.cache_data
def parsear_documento_metodologias(path):
    """
    Analiza un documento .docx y extrae las metodologías basadas en una estructura
    de títulos numéricos (ej. 3.1, 3.1.1, 3.2).
    """
    # Si no existe localmente, intentar descargar desde GitHub
    if not os.path.exists(path):
        if "Ingeniería de Prompts" in path:
            github_url = GITHUB_PROMPTS_URL
            filename = "Ingeniería de Prompts.docx"
        else:
            return {"error": f"No se pudo encontrar el archivo en la ruta: {path}"}
        
        # Descargar desde GitHub
        doc_content = descargar_documento_desde_github(github_url, filename)
        if doc_content is None:
            return {"error": f"No se pudo descargar {filename} desde GitHub."}
        
        try:
            document = docx.Document(doc_content)
        except Exception as e:
            return {"error": f"Ocurrió un problema al leer el archivo `{filename}`: {e}"}
    else:
        # Si existe localmente, leerlo directamente
        try:
            document = docx.Document(path)
        except Exception as e:
            return {"error": f"Ocurrió un problema al leer el archivo `{path}`: {e}"}
    
    try:
        # Inicializar variables para el análisis
        metodologias = {}
        metodologia_activa = None
        subseccion_activa = None 
        
        # Regex más estricta para evitar falsos positivos. Busca el número y el texto.
        titulo_principal_re = re.compile(r"^\s*(3\.\d{1,2})\.\s+(.*)")
        titulo_subseccion_re = re.compile(r"^\s*(3\.\d{1,2}\.\d+)\.\s+(.*)")

        for block in document.element.body:
            if isinstance(block, docx.oxml.text.paragraph.CT_P):
                para_text = docx.text.paragraph.Paragraph(block, document).text.strip()
                if not para_text:
                    continue

                if titulo_principal_re.match(para_text):
                    # Solo actualiza el título activo, no crea la entrada aún.
                    metodologia_activa = para_text 
                    subseccion_activa = None
                elif titulo_subseccion_re.match(para_text):
                    if metodologia_activa:
                        # Ahora, al encontrar la primera subsección, crea la entrada para la metodología.
                        if metodologia_activa not in metodologias:
                            metodologias[metodologia_activa] = {}
                        subseccion_activa = para_text
                        metodologias[metodologia_activa][subseccion_activa] = ""
                elif metodologia_activa and subseccion_activa:
                    metodologias[metodologia_activa][subseccion_activa] += para_text + "\n\n"
            
            elif isinstance(block, docx.oxml.table.CT_Tbl):
                if metodologia_activa and subseccion_activa:
                    # Asegura que la metodología exista antes de añadirle una tabla.
                    if metodologia_activa not in metodologias:
                        metodologias[metodologia_activa] = {}
                    # Si encontramos una tabla, la convertimos a formato Markdown
                    table = docx.table.Table(block, document)
                    
                    # Header
                    header_cells = table.rows[0].cells
                    header = ' | '.join(cell.text.strip() for cell in header_cells)
                    md_table = f"| {header} |\n"
                    
                    # Separator
                    separator = ' | '.join(['---'] * len(header_cells))
                    md_table += f"| {separator} |\n"
                    
                    # Body
                    for row in table.rows[1:]:
                        body_cells = ' | '.join(cell.text.strip() for cell in row.cells)
                        md_table += f"| {body_cells} |\n"
                    
                    metodologias[metodologia_activa][subseccion_activa] += md_table + "\n"

        return metodologias
    except Exception as e:
        return {"error": f"Ocurrió un problema al analizar el archivo: {e}"}

@st.cache_resource
def get_retriever():
    """
    Carga, procesa y vectoriza los documentos para crear un retriever.
    Esta función se cachea para que solo se ejecute una vez.
    """
    try:
        # 1. Cargar y combinar los documentos
        texto_prompts = leer_documento_docx(PROMPTS_DOC_PATH)
        texto_manual = leer_documento_docx(MANUAL_DOC_PATH)
        
        if "**Error:**" in texto_prompts or "**Error:**" in texto_manual:
            st.error("No se pudieron cargar los documentos base para el asistente. Verifica las rutas.")
            return None

        texto_completo = texto_prompts + "\n\n" + texto_manual

        # 2. Dividir el texto en fragmentos
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
        fragmentos = text_splitter.split_text(texto_completo)

        # 3. Configurar embeddings con fallback (Mistral -> Gemini)
        embeddings = None
        if st.secrets.get("MISTRAL_API_KEY"):
            try:
                embeddings = MistralAIEmbeddings(mistral_api_key=st.secrets["MISTRAL_API_KEY"])
            except Exception as e:
                st.warning(f"Error con embeddings de Mistral: {e}. Intentando con Gemini.")
        if not embeddings and st.secrets.get("GEMINI_API_KEY"):
            try:
                embeddings = GoogleGenerativeAIEmbeddings(google_api_key=st.secrets["GEMINI_API_KEY"], model="models/embedding-001")
            except Exception as e:
                st.warning(f"Error con embeddings de Gemini: {e}.")
        
        if not embeddings:
            st.error("No se pudo inicializar ningún modelo de embeddings. Verifica tus API keys de Mistral o Gemini.")
            return None
        
        vectorstore = FAISS.from_texts(texts=fragmentos, embedding=embeddings)
        return vectorstore.as_retriever()

    except Exception as e:
        st.error(f"Error crítico al crear el retriever: {e}")
        return None

def get_rag_chain(retriever, preferred_provider):
    """
    Crea la cadena RAG con el LLM seleccionado y el retriever cacheado.
    """
    try:
        # Configurar LLM con selección de usuario y fallback
        llm = None
        provider_activo = None

        # 4. Configurar LLM con selección de usuario y fallback
        llm = None
        provider_activo = None
        
        # Orden de intento: el preferido primero, luego el resto.
        all_providers = ["Groq", "XAI", "Mistral", "Gemini", "DeepSeek", "OpenAI"]
        # Crea una nueva lista ordenada con el proveedor preferido al principio.
        # Esto es más limpio que modificar la lista original y elimina la advertencia.
        fallback_order = [preferred_provider] + [p for p in all_providers if p != preferred_provider]

        for provider in fallback_order:
            if llm: break # Si ya tenemos un LLM, salimos del bucle
            try:
                if provider == "Groq" and st.secrets.get("GROQ_API_KEY"):
                    llm = ChatGroq(groq_api_key=st.secrets["GROQ_API_KEY"], model_name="llama-3.3-70b-versatile")
                    provider_activo = "Groq (Llama3.3-70b)"
                elif provider == "XAI" and st.secrets.get("XAI_API_KEY"):
                    llm = ChatOpenAI(api_key=st.secrets["XAI_API_KEY"], model="grok", base_url="https://api.x.ai/v1")
                    provider_activo = "XAI (Grok)"
                elif provider == "Mistral" and st.secrets.get("MISTRAL_API_KEY"):
                    llm = ChatMistralAI(mistral_api_key=st.secrets["MISTRAL_API_KEY"], model="mistral-large-latest")
                    provider_activo = "Mistral"
                elif provider == "Gemini" and st.secrets.get("GEMINI_API_KEY"):
                    llm = ChatGoogleGenerativeAI(google_api_key=st.secrets["GEMINI_API_KEY"], model="gemini-1.5-flash", temperature=0.7)
                    provider_activo = "Gemini (Capa Gratuita)"
                elif provider == "DeepSeek" and st.secrets.get("DEEPSEEK_API_KEY"):
                    # Usar requests directo como en tu proyecto CAMACOL
                    def deepseek_invoke(messages):
                        url = "https://api.deepseek.com/v1/chat/completions"
                        headers = {
                            'Content-Type': 'application/json',
                            'Authorization': f"Bearer {st.secrets['DEEPSEEK_API_KEY']}"
                        }
                        payload = {
                            "model": "deepseek-chat",
                            "messages": messages,
                            "temperature": 0.7,
                            "max_tokens": 2000
                        }
                        response = requests.post(url, headers=headers, json=payload, timeout=30)
                        if response.status_code == 200:
                            data = response.json()
                            return data['choices'][0]['message']['content']
                        else:
                            raise Exception(f"DeepSeek API Error: {response.status_code} - {response.text}")
                    
                    # Crear un wrapper compatible con LangChain
                    from langchain_core.language_models import BaseLanguageModel
                    from langchain_core.messages import BaseMessage, HumanMessage
                    from langchain_core.outputs import LLMResult
                    from typing import List, Optional, Any, Dict
                    
                    class DeepSeekLLM(BaseLanguageModel):
                        def __init__(self, invoke_func):
                            super().__init__()
                            self.invoke_func = invoke_func
                        
                        def _generate(self, messages: List[BaseMessage], stop: Optional[List[str]] = None, run_manager: Optional[Any] = None, **kwargs: Any) -> LLMResult:
                            # Convertir mensajes a formato de la API
                            api_messages = []
                            for msg in messages:
                                if isinstance(msg, HumanMessage):
                                    api_messages.append({"role": "user", "content": msg.content})
                                else:
                                    api_messages.append({"role": "assistant", "content": msg.content})
                            
                            response = self.invoke_func(api_messages)
                            
                            # Crear generación
                            from langchain_core.outputs import Generation
                            generation = Generation(text=response)
                            return LLMResult(generations=[[generation]])
                        
                        def _llm_type(self) -> str:
                            return "deepseek"
                        
                        def invoke(self, input: Any, config: Optional[Dict[str, Any]] = None, **kwargs: Any) -> str:
                            if isinstance(input, str):
                                return self.invoke_func([{"role": "user", "content": input}])
                            elif isinstance(input, list) and input and isinstance(input[0], BaseMessage):
                                api_messages = []
                                for msg in input:
                                    if isinstance(msg, HumanMessage):
                                        api_messages.append({"role": "user", "content": msg.content})
                                    else:
                                        api_messages.append({"role": "assistant", "content": msg.content})
                                return self.invoke_func(api_messages)
                            else:
                                return str(input)
                        
                        async def ainvoke(self, input: Any, config: Optional[Dict[str, Any]] = None, **kwargs: Any) -> str:
                            return self.invoke(input, config, **kwargs)
                    
                    llm = DeepSeekLLM(deepseek_invoke)
                    provider_activo = "DeepSeek"
                elif provider == "OpenAI" and st.secrets.get("OPENAI_API_KEY"):
                    llm = ChatOpenAI(api_key=st.secrets["OPENAI_API_KEY"], model="gpt-4o-mini", temperature=0.7, base_url="https://api.openai.com/v1")
                    provider_activo = "OpenAI (GPT-4o-mini)"
            except Exception as e:
                st.warning(f"No se pudo inicializar el proveedor '{provider}': {e}. Intentando con el siguiente.")
                continue

        if not llm:
            st.error("No se pudo inicializar ningún LLM. Verifica tus API keys en el archivo secrets.toml.")
            return None

        # 5. Crear la cadena conversacional con LCEL (LangChain Expression Language)
        prompt_template = (
            "Eres un asistente experto en ingeniería de prompts para AGV Open Market. "
            "Usa el siguiente contexto para responder la pregunta. "
            "Tu objetivo es guiar al usuario de negocio para construir el prompt perfecto. "
            "Recomienda la metodología más adecuada (ej. Zero-Shot, Chain of Thought), la plataforma (Copilot 365, Copilot Studio), y proporciona el prompt final. "
            "Si el usuario pide un diagrama, ilustración o imagen, incluye en tu respuesta el tag [IMAGE: descripción detallada de la imagen en inglés]. "
            "Si no sabes la respuesta, di que no tienes suficiente información en tus documentos. "
            "Sé conciso y directo."
            "\n\n"
            "Contexto: {context}\n"
            "Pregunta: {question}\n"
            "Respuesta:"
        )
        prompt = ChatPromptTemplate.from_template(prompt_template)

        # Construcción de la cadena con LCEL
        rag_chain = (
            {"context": retriever, "question": RunnablePassthrough()}
            | prompt
            | llm
            | StrOutputParser()
        )
        return rag_chain, provider_activo

    except Exception as e:
        st.error(f"Error al construir la cadena RAG: {e}")
        return None

@st.cache_data
def generate_image(prompt):
    """Genera una imagen usando Vertex AI y devuelve la imagen o None."""
    try:
        # Extraer credenciales de los secretos de Streamlit
        project_id = st.secrets.get("VERTEX_PROJECT_ID")
        location = st.secrets.get("VERTEX_LOCATION", "us-central1")

        if not project_id:
            st.warning("El ID del proyecto de Vertex AI no está configurado en los secretos. No se pueden generar imágenes.")
            return None

        vertexai.init(project=project_id, location=location)
        model = ImageGenerationModel.from_pretrained("imagegeneration@006")
        images = model.generate_images(prompt=prompt, number_of_images=1)
        return images[0]._image_bytes
    except Exception as e:
        st.error(f"Error al generar la imagen: {e}")
        return None

def load_local_css():
    """Carga CSS personalizado para mejorar la estética de la aplicación."""
    st.markdown("""
        <style>
        /* Estilo para los expanders que contienen la información de la metodología */
        [data-testid="stExpander"] {
            border: 1px solid #dee2e6;
            border-radius: 10px;
            box-shadow: 2px 2px 12px rgba(0,0,0,0.05);
            margin-bottom: 1rem;
        }
        /* Estilo para el encabezado del expander */
        [data-testid="stExpander"] > div:first-child {
            background-color: #f8f9fa;
        }
        </style>
    """, unsafe_allow_html=True)

def mostrar_infografia_dinamica(titulo_metodologia, contenido):
    """Muestra una 'infografía' con títulos de subsección limpios y re-numerados."""
    st.subheader(titulo_metodologia, divider="rainbow")

    if not contenido:
        st.warning("No se encontró contenido detallado para esta metodología.")
        return
    
    # Ordena las subsecciones para asegurar una renumeración consistente
    def natural_sort_key(s):
        return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]
    
    sorted_subsecciones = sorted(contenido.keys(), key=natural_sort_key)

    for i, titulo_original in enumerate(sorted_subsecciones):
        texto_subseccion = contenido[titulo_original]
        # Elimina el prefijo numérico original (ej. "3.1.1. ") y añade el nuevo
        titulo_limpio = f"{i + 1}. {re.sub(r'^\s*\d+(\.\d+)*\.\s*', '', titulo_original)}"
        with st.expander(titulo_limpio, expanded=True):
            st.markdown(texto_subseccion)

# --- Interfaz de Usuario ---
load_local_css()

col1, col2 = st.columns([1, 4])
with col1:
    if os.path.exists(LOGO_PATH):
        st.image(LOGO_PATH, width=150)
    else:
        st.warning("Logo no encontrado.")

with col2:
    st.title("Visor de Documentos de AGV Open Market")
    st.write("Una herramienta para explorar la documentación de Ingeniería de Prompts y manuales de agentes.")

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["Prompt Efectivo", "Metodologías", "System Prompts", "Conocimiento", "Ejemplos para el Asistente", "Asistente de Prompts"])

with tab2:
    st.header("Análisis Dinámico del Documento de Metodologías")

    with st.spinner("Analizando `Ingeniería de Prompts.docx`..."):
        datos_metodologias = parsear_documento_metodologias(PROMPTS_DOC_PATH)

    if "error" in datos_metodologias:
        st.error(datos_metodologias["error"])
    elif not datos_metodologias:
        st.warning("No se pudo encontrar ninguna metodología con el formato esperado (ej. '3.1.', '3.1.1.') en el documento.")
    else:
        # Ordena las llaves (títulos) usando un ordenamiento natural para que 3.10 vaya después de 3.9
        def natural_sort_key(s):
            # Extrae los números de la cadena y los convierte a enteros para un ordenamiento correcto
            return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

        lista_original = sorted(datos_metodologias.keys(), key=natural_sort_key)
        
        # Crea un mapeo de títulos limpios a títulos originales
        # ej: "1. Zero-Shot..." -> "3.1. Zero-Shot..."
        mapa_titulos = {f"{i+1}.{''.join(title.split('.')[1:]).lstrip()}": title for i, title in enumerate(lista_original)}
        lista_limpia = list(mapa_titulos.keys())
        
        # Menú de selección para elegir qué metodología ver
        opcion_limpia_seleccionada = st.selectbox(
            "Selecciona una metodología para ver su detalle:",
            lista_limpia
        )

        if opcion_limpia_seleccionada:
            # Usa el mapa para encontrar el título original y obtener los datos
            titulo_original = mapa_titulos[opcion_limpia_seleccionada]
            # Muestra la "infografía" para la metodología seleccionada
            mostrar_infografia_dinamica(opcion_limpia_seleccionada, datos_metodologias[titulo_original])

with tab1:
    st.header("Anatomía y Estructura de un Prompt Efectivo")

    with st.expander("1. Los 5 Pilares de un Prompt Efectivo (Texto)", expanded=True):
        st.markdown("""
        Para las interacciones estándar con la IA generativa, la fórmula fundamental se compone de cinco elementos clave:
        1.  **Rol:** Define la personalidad o perfil profesional que debe adoptar la IA (ej. "Eres un analista de mercado experto" o "redactor publicitario").
        2.  **Tarea o Instrucción:** Indica claramente la acción concreta a realizar mediante verbos de acción (ej. "resume", "verifica", "crea un ticket", "traduce").
        3.  **Contexto:** Proporciona los antecedentes necesarios para que la IA entienda el entorno de la tarea (ej. "basado en el informe de ventas del Q3" o "para un público no técnico").
        4.  **Ejemplos (Few-shot):** Incluir uno o varios ejemplos de la dinámica "Pregunta -> Respuesta esperada" ayuda enormemente a la IA a comprender matices complejos y el formato deseado.
        5.  **Formato y Tono:** Define cómo debe presentarse la información (ej. "en una tabla comparativa", "en párrafos cortos", "en formato Markdown") y la actitud profesional (ej. "tono empático", "formal", "directo").
        """)

    with st.expander("2. Anatomía de Prompts de Sistema (System Prompts para Agentes)", expanded=True):
        st.markdown("""
        Cuando se diseña el "cerebro" de un agente en Copilot 365, la estructura debe ser más robusta y jerarquizada debido al límite técnico de 8,000 caracteres.
        - **Uso de Markdown:** Se deben emplear títulos (`#` para propósitos, `##` para pasos) y negritas para resaltar restricciones críticas, facilitando que el orquestador de la IA priorice las reglas de seguridad sobre las tareas generales.
        - **Referencia Explícita a Fuentes:** El prompt debe guiar al agente sobre cuándo y cómo usar cada archivo conectado (ej. "Usa Inventario_Q1.xlsx solo para preguntas de stock").
        - **Guardrails (Reglas de Oro):** Es vital incluir instrucciones para evitar alucinaciones, ordenando a la IA que si no encuentra evidencia en los documentos, admita que no sabe la respuesta en lugar de inventarla.
        - **Responsabilidad y Atribución:** El prompt de sistema debe forzar al agente a presentarse como un "empleado aumentado" y a incluir siempre una nota de que el contenido fue asistido por IA y verificado por un humano.
        """)

    with st.expander("3. Estructura de Prompts Multimedia (Imagen y Vídeo)", expanded=True):
        st.markdown("""
        Para la generación de contenido visual, la anatomía del prompt cambia de una estructura narrativa a una fórmula técnica:
        - **Fórmula Visual:** Sujeto principal + tipo de plano (ej. primer plano, plano medio) + detalles del objeto + detalles del fondo + iluminación y atmósfera.
        - **Prompts de Vídeo:** A la fórmula anterior se le añade obligatoriamente el movimiento de cámara (ej. plano dolly, paneo rápido, plano de seguimiento) para dirigir la animación.
        - **Idioma:** Se recomienda redactar estos prompts técnicos en inglés para maximizar la precisión de la herramienta, independientemente del idioma del usuario.
        """)
        st.markdown("""
        | Componente | Función Técnica | Impacto en el Resultado |
        | :--- | :--- | :--- |
        | Verbos concretos | Define la operación. | Evita respuestas vagas o incompletas. |
        | Markdown (#) | Jerarquiza el razonamiento. | Asegura que se respeten primero las reglas de seguridad. |
        | Citas/Grounded | Ancla la IA a datos reales. | Elimina alucinaciones al forzar el uso de documentos oficiales. |
        | Few-shot | Proporciona un patrón visual. | Mejora drásticamente la consistencia en tareas repetitivas. |
        """)

with tab3:
    st.header("System Prompts")
    st.markdown("""
    Esta metodología convierte a los documentos estáticos de la empresa en instrucciones dinámicas, permitiendo que el agente evolucione a la par de la identidad visual y narrativa de la organización.
    """)

    with st.expander("Configuración del 'cerebro' del agente", expanded=True):
        st.markdown("""
        La configuración del "cerebro" del agente, conocida técnicamente como System Prompt o Instrucciones, es el pilar fundamental que define cómo Microsoft 365 Copilot procesará la información y se comportará ante el usuario. Estas instrucciones actúan como un "Prompt Coach" invisible, optimizando la calidad de cada respuesta incluso antes de que el usuario final realice su primera consulta.
        """)
        st.info("""
        **El Límite Técnico: Los 8,000 Caracteres**
        En plataformas como Copilot 365, existe un límite estricto de 8,000 caracteres para estas instrucciones.
        - **Gestión del espacio:** Este límite obliga al "Prompt Engineer" a ser sumamente preciso, eliminando verbosidad innecesaria y centrándose en comandos directos.
        - **Priorización:** Dado el espacio finito, es vital utilizar estructuras como la metodología MoSCoW para asegurar que las reglas críticas de seguridad y veracidad ocupuen los primeros caracteres, garantizando que el modelo las priorice sobre aspectos cosméticos.
        """)
        st.info("""
        **Definición del Rol: La Personalidad Profesional**
        Establecer una personalidad clara es el primer paso para evitar respuestas genéricas y alinear al asistente con la organización.
        - **Identidad:** Se debe especificar quién es el agente (ej. "Eres un analista financiero experto en el mercado latinoamericano").
        - **Contexto de actuación:** Al definir el rol, se delimita el campo de acción del agente, lo que reduce la probabilidad de que responda sobre temas fuera de su competencia.
        """)
        st.info("""
        **Componentes Esenciales de la Instrucción**
        Para construir un cerebro robusto, las instrucciones deben cubrir cuatro áreas críticas:
        - **Objetivo (Propósito):** Declarar qué debe lograr exactamente el agente (ej. "Tu propósito es ayudar a los empleados a interpretar las políticas de gastos").
        - **Pasos Lógicos (Flujo de trabajo):** Enumerar la secuencia de acciones que la IA debe realizar internamente antes de responder. Por ejemplo, "1. Busca en SharePoint, 2. Resume los hallazgos, 3. Verifica si hay datos sensibles".
        - **Tono y Estilo:** Determinar la actitud comunicativa. Las fuentes sugieren tonos que van desde un "profesor amable y paciente" hasta un estilo formal y conciso.
        - **Fuentes de Conocimiento Disponibles:** Es imperativo hacer una referencia explícita a los datos conectados (grounding), indicando cuándo y cómo usar cada recurso. Por ejemplo, se debe instruir al agente para que use la "Carpeta de Manuales" solo cuando se pidan procedimientos detallados.
        """)
        st.info("""
        **Ventajas de una Configuración Detallada**
        - **Consistencia Corporativa:** Garantiza que todos los usuarios reciban respuestas bajo los mismos estándares de calidad y tono, independientemente de la complejidad de su pregunta.
        - **Reducción de Alucinaciones:** Al limitar explícitamente el conocimiento al contenido de la empresa y dar instrucciones de veracidad, el agente se vuelve un colaborador mucho más fiable que la IA pública.
        - **Automatización de Tareas:** Permite que el agente no solo responda preguntas, sino que actúe sobre flujos de trabajo (ej. crear tickets en Planner) mediante el uso de verbos de acción concretos como "crea", "verifica" o "notifica".
        """)

    with st.expander("Uso de Markdown para jerarquizar reglas y priorizar fuentes", expanded=True):
        st.markdown("""
        El uso de Markdown en las instrucciones de sistema no es una cuestión estética, sino una herramienta técnica fundamental para estructurar el razonamiento del modelo de IA y asegurar que priorice las reglas críticas de negocio. Al utilizar esta sintaxis, el Prompt Engineer proporciona una señal clara al orquestador sobre la jerarquía de la información, lo que reduce la carga cognitiva del modelo y mejora la precisión en la ejecución de tareas complejas.
        - **Estructura Jerárquica:** Se deben utilizar encabezados que segreguen las funciones del agente (`#` para Rol, `##` para Pasos, `###` para excepciones).
        - **Resaltado de Instrucciones:** El uso de **negritas** para restricciones críticas y `monoespaciado` para palabras clave ayuda al modelo a identificar qué partes del prompt son innegociables.
        - **Referencia Explícita a Fuentes:** Es imperativo guiar al agente sobre cuándo y cómo debe interactuar con cada archivo conectado (ej. "Usa `Inventario_Tecnico.xlsx` solo para preguntas de stock").
        - **Ciclo de Refinamiento Iterativo:** La implementación de Markdown debe ser probada en el Test Canvas. Si el agente ignora una regla, se debe elevar su jerarquía en el Markdown hasta que el comportamiento sea consistente.
        """)

    with st.expander("Definición de Guardrails (Reglas de Oro) para evitar datos inventados", expanded=True):
        st.markdown("""
        Los Guardrails son las fronteras de seguridad que impiden que el agente "alucine" o comprometa la integridad de la información corporativa.
        - **Regla de Silencio:** Se debe incluir una instrucción clara: "Si la información no se encuentra en los documentos conectados, admite que no la conoces y no intentes inventar una respuesta".
        - **Uso de Verbos de Acción:** Para reducir la ambigüedad, el prompt debe usar comandos concretos como "verifica", "busca", "pregunta" o "cita", en lugar de instrucciones vagas.
        - **Validación de Citación:** Obligar al agente a incluir enlaces directos a los documentos originales de SharePoint o OneDrive asegura la trazabilidad y permite que el usuario humano valide el resultado.
        - **Manejo de Limitaciones:** Definir explícitamente qué no debe hacer el agente (ej. "No exponer secretos comerciales" o "No procesar datos PII") es fundamental para el cumplimiento del RGPD y la ética de IA responsable.
        """)

with tab4:
    st.header("Gestión de Fuentes y Conocimiento")

    with st.expander("1. Referencia explícita a archivos: Instruir cuándo usar SharePoint vs. OneDrive", expanded=True):
        st.markdown("""
        La distinción entre SharePoint y OneDrive no es solo de almacenamiento, sino de contexto operativo y jerarquía de la información. La Ingenieria de prompts debe instruir al agente sobre qué "biblioteca" consultar según el tipo de solicitud del usuario.
        - **SharePoint (Conocimiento Departamental/Oficial):** Se debe priorizar para manuales de procesos, políticas vigentes, wikis de ingeniería y documentación técnica compartida. Al configurar el prompt de sistema, se debe especificar que SharePoint es la "fuente de verdad" para consultas normativas o de equipo.
        - **OneDrive (Conocimiento Individual/Específico):** Se utiliza para conectar documentos personales, archivos de trabajo en progreso o currículums específicos que no pertenecen a una biblioteca pública. Es útil cuando el agente asiste a un solo usuario en tareas de redacción o análisis de sus propios borradores.
        - **Instrucciones Técnicas de Selección:** En el campo de instrucciones, se debe escribir explícitamente el uso de cada origen; por ejemplo, instruir al agente para que use la carpeta `Manual_Procesos_2025` de SharePoint solo cuando se pidan pasos detallados.
        - **Límites de los Archivos:** Es vital recordar que los archivos cargados tienen un límite de 512 MB y deben ser formatos compatibles como PDF, DOCX o XLSX. SharePoint permite una búsqueda mejorada con archivos de hasta 200 MB si se activa la función Work IQ.
        """)

    with st.expander("2. Prompts de citación obligatoria: Forzar la veracidad mediante enlaces directos", expanded=True):
        st.markdown("""
        Para eliminar el riesgo de alucinaciones y fomentar la confianza del usuario, el diseño del prompt debe incluir una directiva de citación obligatoria. Esto obliga a la IA a demostrar de dónde extrajo la información.
        - **Regla de Atribución:** El prompt de sistema debe exigir que cada respuesta incluya el nombre del documento y un enlace directo al archivo original en SharePoint o OneDrive. Esto permite que el usuario valide el contenido de forma manual inmediatamente.
        - **Citas a Nivel de Página:** Para documentos extensos en formato PDF, Microsoft 365 Copilot permite citas a nivel de página. La Ingenieria de prompts debe instruir al agente: "Si utilizas un PDF, indica el número de página de donde extrajiste la conclusión".
        - **Manejo de Silencio (Guardrails):** Se debe incluir una restricción crítica: "Si la respuesta no se encuentra en los archivos citados, admite que no lo sabes en lugar de intentar adivinar".
        - **Ejemplo de Prompt de Citación:** *"Eres un analista legal. Responde a las dudas del usuario basándote exclusivamente en la carpeta 'Contratos_2024'. OBLIGACIÓN: Al final de cada respuesta, añade una sección titulada 'Fuentes Consultadas' con el enlace directo al archivo de SharePoint utilizado."*
        """)

with tab5:
    st.header("Ejemplos de Preguntas para el Asistente")
    st.info("Usa estas preguntas como inspiración en la pestaña 'Asistente de Prompts'. El asistente utilizará el conocimiento de los manuales para darte respuestas detalladas y personalizadas.")

    # Categorías de preguntas
    categorias = {
        "1. Construcción y Diseño de Prompts": [
            "Necesito un prompt para resumir un informe de ventas trimestral para la junta directiva.",
            "¿Cómo escribo un prompt para clasificar correos de soporte en 'Urgente', 'Consulta' y 'Feedback'?",
            "Dame un prompt para extraer todas las fechas, montos y partes involucradas de un contrato en PDF.",
            "Quiero un prompt que genere 5 ideas creativas para una campaña de marketing de un nuevo producto.",
            "¿Cuál es la mejor forma de pedirle a la IA que compare dos productos basándose en sus especificaciones técnicas?",
            "Necesito un prompt que actúe como un 'abogado del diablo' para encontrar fallas en mi plan de proyecto.",
            "Dame un prompt para generar un guion para un video de capacitación de 2 minutos sobre seguridad informática.",
            "¿Cómo le pido a la IA que reescriba un texto técnico para una audiencia no especializada?",
            "Quiero un prompt que genere una tabla Markdown comparando las características de tres software diferentes.",
            "Necesito un prompt que simule una entrevista de trabajo para un puesto de 'Jefe de Producto'."
        ],
        "2. Metodologías de Prompting": [
            "Tengo un proceso de aprobación de facturas con varios pasos. ¿Qué metodología de prompt debo usar?",
            "Explícame la diferencia entre Zero-Shot y Few-Shot Prompting con un ejemplo de negocio.",
            "¿En qué escenario es más útil la metodología MoSCoW para estructurar un prompt?",
            "¿Cómo puedo usar 'Self-Reflection' para que un prompt mejore su propia respuesta?",
            "Dame un ejemplo de 'Prompting Estructural' usando una plantilla para generar informes de incidentes.",
            "¿Qué es la 'Orquestación Multi-Agente' y cómo podría aplicarla para un proceso de onboarding de nuevos empleados?",
            "Si tengo muchos ejemplos, ¿es mejor usar Few-Shot en el prompt o Few-Shot por archivo?",
            "¿Para qué tipo de problema matemático es ideal el 'Chain of Thought'?",
            "¿Cómo funciona el 'Prompt Coaching' en una conversación para refinar una idea?",
            "Quiero generar imágenes de productos. ¿Qué metodología de prompting multimedia debo seguir?"
        ],
        "3. Creación de Agentes (Copilot 365 y Copilot Studio)": [
            "Quiero crear mi primer agente en Copilot 365 para responder preguntas sobre las políticas de vacaciones. ¿Cuáles son los pasos?",
            "¿Cuándo debo usar Copilot Studio en lugar del Agent Builder de SharePoint? Dame un árbol de decisión.",
            "¿Qué KPIs debo usar para medir el éxito de un agente que gestiona tickets de TI?",
            "¿Cómo defino el 'Mapa de Conocimiento' para un agente que debe consultar manuales de productos y listas de precios?",
            "Explícame la diferencia entre un agente declarativo y uno con motor personalizado en términos de esfuerzo y capacidad.",
            "¿Puedo crear un agente que entienda la jerga específica de mi empresa? ¿Cómo lo entreno?",
            "¿Cuál es el proceso para publicar un agente creado en Copilot Studio para que esté disponible en Microsoft Teams?",
            "¿Cómo puedo personalizar el ícono y el mensaje de bienvenida de mi agente para que coincida con la marca de la empresa?",
            "¿Qué es el 'ciclo de refinamiento iterativo' y cómo lo aplico para mejorar mi agente?",
            "Mi agente da respuestas incorrectas. ¿Cómo uso el panel de prueba para depurarlo?"
        ],
        "4. System Prompts y Guardrails de Seguridad": [
            "Dame un ejemplo de un 'System Prompt' robusto para un agente de RRHH.",
            "¿Cómo uso Markdown en las instrucciones para que el agente priorice las reglas de seguridad?",
            "¿Qué es una 'Regla de Silencio' y cómo la escribo en el System Prompt para evitar alucinaciones?",
            "Necesito una regla para que mi agente nunca revele información de salarios. ¿Cómo la formulo?",
            "Explícame cómo las negritas y las listas numeradas afectan el comportamiento del agente en el System Prompt.",
            "¿Por qué es importante el límite de 8,000 caracteres en las instrucciones y cómo lo gestiono?",
            "Dame un ejemplo de una instrucción para forzar al agente a citar sus fuentes.",
            "¿Qué significa 'definir explícitamente qué no debe hacer el agente' y por qué es crucial?",
            "¿Cómo le indico a mi agente que debe usar un tono 'amable pero profesional'?",
            "¿Cuál es la diferencia entre el 'Rol' y el 'Propósito' en un System Prompt?"
        ],
        "5. Gestión de Conocimiento y Automatización": [
            "Mi agente necesita enviar un correo de notificación. ¿Cómo lo integro con Power Automate?",
            "¿Cómo le digo a mi agente que use un archivo de SharePoint para políticas y uno de OneDrive para borradores?",
            "Explícame cómo un agente puede usar un conector para acceder a datos de una API externa.",
            "¿Qué es Microsoft Graph y cómo ayuda a que el agente respete los permisos de los usuarios?",
            "Tengo un manual de 300 páginas en PDF. ¿Cómo le pido al agente que cite el número de página en sus respuestas?",
            "¿Cuáles son los límites de tamaño y formato de archivo para el conocimiento de un agente?",
            "¿Qué son las 'Acciones' en Copilot Studio y cómo le dan 'músculos' a mi agente?",
            "Dame un ejemplo de un flujo de Power Automate que un agente podría disparar para crear una tarea en Planner.",
            "¿Cómo me aseguro de que mi agente solo use la 'fuente de verdad' oficial y no mezcle información?",
            "¿Qué es 'grounding' y por qué es la base para evitar alucinaciones?"
        ],
        "6. Optimización y Resolución de Problemas": [
            "Mi agente es muy lento. ¿Qué puedo hacer para optimizar su rendimiento?",
            "Las respuestas de mi agente son demasiado genéricas. ¿Cómo puedo hacerlas más específicas?",
            "El agente a veces ignora mis instrucciones de formato. ¿Cómo puedo reforzar el formato de salida?",
            "¿Qué hago si el agente 'alucina' y se inventa datos a pesar de mis guardrails?",
            "¿Cómo puedo usar el 'Test Canvas' para probar escenarios límite y excepciones?",
            "El agente no encuentra información en un documento que sé que está conectado. ¿Cuál podría ser el problema?",
            "¿Cómo puedo ver qué fuentes está usando el agente para construir su respuesta?",
            "Mi prompt es muy largo. ¿Cómo puedo hacerlo más conciso sin perder efectividad?",
            "¿Es mejor darle al agente muchos documentos pequeños o un solo documento grande?",
            "¿Cómo puedo monitorear el uso y la satisfacción de los usuarios con mi agente una vez publicado?"
        ]
    }

    for categoria, preguntas in categorias.items():
        with st.expander(categoria):
            for pregunta in preguntas:
                st.markdown(f"- **Pregunta:** \"{pregunta}\"")

with tab6:
    st.header("Asistente Interactivo para la Creación de Prompts")
    st.markdown("Chatea con un asistente experto para diseñar el prompt perfecto para tu necesidad de negocio.")

    # Opciones de proveedores de LLM
    provider_options = ["Groq", "XAI", "Mistral", "Gemini", "DeepSeek", "OpenAI"]
    selected_provider = st.selectbox("Selecciona tu motor de IA preferido:", provider_options)

    # Cargar el retriever una sola vez y cachearlo
    retriever = get_retriever()

    if retriever:
        # Inicializa el asistente y el historial de chat en la sesión
        session_key = f"conv_chain_{selected_provider}"
        if st.session_state.get("current_session_key") != session_key:
            with st.spinner("Configurando el asistente con el conocimiento de AGV..."):
                result = get_rag_chain(retriever, selected_provider)
                if result:
                    st.session_state.conversation_chain, st.session_state.provider_activo = result
                    st.session_state.current_session_key = session_key
                else:
                    st.session_state.conversation_chain = None
        
        if "chat_history" not in st.session_state or st.session_state.get("current_session_key") != session_key:
            st.session_state.chat_history = []

        if st.session_state.conversation_chain:
            # Muestra el historial del chat
            for message in st.session_state.chat_history:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            # Input del usuario
            if prompt := st.chat_input("¿En qué proceso de negocio necesitas ayuda?"):
                st.session_state.chat_history.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)

                with st.chat_message("assistant"):
                    with st.spinner("Pensando..."):
                        response_text = st.session_state.conversation_chain.invoke(prompt)
                        
                        # Procesar para generar imágenes si se solicita
                        image_tag_match = re.search(r"\[IMAGE:\s*(.*?)\]", response_text)
                        if image_tag_match and st.secrets.get("GEMINI_API_KEY"):
                            image_prompt = image_tag_match.group(1)
                            # Limpiar el tag de la respuesta de texto
                            response_text = re.sub(r"\[IMAGE:\s*(.*?)\]", "", response_text).strip()
                            st.markdown(f"**Respuesta de {st.session_state.provider_activo}:**\n\n{response_text}")
                            with st.spinner("Generando imagen..."):
                                image_bytes = generate_image(image_prompt)
                                if image_bytes:
                                    st.image(image_bytes, caption=f"Ilustración para: '{image_prompt}'")
                        else:
                            st.markdown(f"**Respuesta de {st.session_state.provider_activo}:**\n\n{response_text}")
                
                st.session_state.chat_history.append({"role": "assistant", "content": response_text})
